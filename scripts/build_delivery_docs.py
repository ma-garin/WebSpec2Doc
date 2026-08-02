#!/usr/bin/env python3
"""SDLC 文書（Markdown 正本）から納品用の Word / Excel を生成する。

Markdown を正本にしているのは、git で差分が追えることと、コードと同じ流れで
レビューできるため。Office ファイルは常にここから再生成する。手で編集すると
正本と乖離し、次のコード変更時に更新されずに陳腐化する。

使い方:
    venv/bin/python scripts/build_delivery_docs.py

出力（正本 Markdown と同じディレクトリに並べる）:
    docs/sdlc/<工程>/WS2D-xx-001_*.docx    全 SDLC 文書
    docs/sdlc/<工程>/WS2D-xx-001_*.xlsx    一覧・マトリクス系（機械抽出データから直接生成）
"""

from __future__ import annotations

import hashlib
import json
import re
import shutil
import subprocess
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

ROOT = Path(__file__).resolve().parent.parent
SDLC_DIR = ROOT / "docs" / "sdlc"
ASBUILT_DIR = SDLC_DIR / "_asbuilt"
FIGURE_CACHE = SDLC_DIR / "_asbuilt" / "figures"

JP_FONT = "Yu Gothic"
MONO_FONT = "Consolas"
DOC_DATE = "2026-08-02"
AUTHOR = "開発チーム"
CONFIDENTIALITY = "社外秘"
PROJECT_NAME = "WebSpec2Doc"

_DIAGRAM_LANGS = {"mermaid"}


# ---------------------------------------------------------------- 図のレンダリング


def _mmdc_command() -> list[str] | None:
    """mermaid-cli の起動コマンド。ローカル導入を優先し、無ければ npx に落とす。"""
    local = ROOT / "node_modules" / ".bin" / "mmdc"
    if local.exists():
        return [str(local)]
    if shutil.which("mmdc"):
        return ["mmdc"]
    if shutil.which("npx"):
        return ["npx", "-y", "@mermaid-js/mermaid-cli@11"]
    return None


def render_mermaid(code: str) -> Path | None:
    """mermaid 定義を PNG にする。内容ハッシュでキャッシュし、再実行を避ける。

    Word はコードブロックを図として描かない。図の定義をそのまま載せた文書は
    「図が無い」のと同じなので、画像に落としてから埋め込む。
    """
    FIGURE_CACHE.mkdir(parents=True, exist_ok=True)
    digest = hashlib.sha256(code.encode("utf-8")).hexdigest()[:16]
    png = FIGURE_CACHE / f"{digest}.png"
    if png.exists() and png.stat().st_size > 0:
        return png

    command = _mmdc_command()
    if command is None:
        return None

    src = FIGURE_CACHE / f"{digest}.mmd"
    src.write_text(code, encoding="utf-8")
    try:
        result = subprocess.run(
            [*command, "-i", str(src), "-o", str(png), "-b", "white", "-w", "1600", "-s", "2"],
            capture_output=True,
            text=True,
            timeout=180,
            check=False,
            cwd=ROOT,
        )
        if result.returncode != 0:
            print(f"WARN: mermaid render failed ({digest}): {result.stderr.strip()[:160]}")
            return None
    except (OSError, subprocess.SubprocessError) as exc:
        print(f"WARN: mermaid render error ({digest}): {exc}")
        return None
    return png if png.exists() else None


# ---------------------------------------------------------------- Markdown 解析


def _split_table_row(line: str) -> list[str]:
    """`| a | b |` を ['a','b'] にする。両端の空セルだけ落とす。"""
    cells = line.split("|")
    if cells and not cells[0].strip():
        cells = cells[1:]
    if cells and not cells[-1].strip():
        cells = cells[:-1]
    return [c.strip() for c in cells]


def _is_table_separator(line: str) -> bool:
    return bool(re.fullmatch(r"\s*\|?[\s:\-|]+\|?\s*", line)) and "-" in line


def parse_markdown(text: str) -> list[dict[str, Any]]:
    """Markdown を Word 生成に必要な最小のブロック列へ落とす。

    pandoc が使えない環境なので自前で解析する。対応するのは見出し・段落・箇条書き・
    番号付き・表・コードブロック・水平線のみ。SDLC 文書はこの範囲で書かれている。
    """
    blocks: list[dict[str, Any]] = []
    lines = text.splitlines()
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # コードブロック
        if stripped.startswith("```"):
            lang = stripped[3:].strip()
            body: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                body.append(lines[i])
                i += 1
            i += 1
            blocks.append({"type": "code", "lang": lang, "text": "\n".join(body)})
            continue

        # 見出し
        if m := re.match(r"^(#{1,6})\s+(.*)$", stripped):
            blocks.append({"type": "heading", "level": len(m.group(1)), "text": m.group(2)})
            i += 1
            continue

        # 水平線
        if re.fullmatch(r"(-{3,}|\*{3,}|_{3,})", stripped):
            blocks.append({"type": "hr"})
            i += 1
            continue

        # 表（次行が区切り行なら表とみなす）
        if stripped.startswith("|") and i + 1 < len(lines) and _is_table_separator(lines[i + 1]):
            header = _split_table_row(stripped)
            i += 2
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(_split_table_row(lines[i].strip()))
                i += 1
            blocks.append({"type": "table", "header": header, "rows": rows})
            continue

        # 箇条書き / 番号付き
        if re.match(r"^\s*([-*+]|\d+\.)\s+", line):
            items: list[dict[str, Any]] = []
            while i < len(lines) and re.match(r"^\s*([-*+]|\d+\.)\s+", lines[i]):
                raw = lines[i]
                indent = (len(raw) - len(raw.lstrip())) // 2
                ordered = bool(re.match(r"^\s*\d+\.", raw))
                content = re.sub(r"^\s*([-*+]|\d+\.)\s+", "", raw)
                items.append({"text": content, "indent": indent, "ordered": ordered})
                i += 1
            blocks.append({"type": "list", "items": items})
            continue

        # 段落（空行まで結合）
        para: list[str] = []
        while i < len(lines) and lines[i].strip() and not _starts_block(lines[i]):
            para.append(lines[i].strip())
            i += 1
        if para:
            blocks.append({"type": "para", "text": " ".join(para)})
        else:
            i += 1

    return blocks


def _starts_block(line: str) -> bool:
    s = line.strip()
    return (
        s.startswith("#")
        or s.startswith("```")
        or s.startswith("|")
        or bool(re.match(r"^\s*([-*+]|\d+\.)\s+", line))
        or bool(re.fullmatch(r"(-{3,}|\*{3,}|_{3,})", s))
    )


# ---------------------------------------------------------------- Word 生成


def _set_jp_font(run, name: str = JP_FONT, size: int | None = None) -> None:
    """日本語フォントは eastAsia 属性も指定しないと游ゴシックにならない。"""
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    if size:
        run.font.size = Pt(size)


def _add_inline(paragraph, text: str) -> None:
    """**強調**・`コード`・[リンク](url) を最低限だけ描き分ける。"""
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1（\2）", text)
    for part in re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            _set_jp_font(run)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            _set_jp_font(run, MONO_FONT)
            run.font.size = Pt(9)
        else:
            run = paragraph.add_run(part)
            _set_jp_font(run)


def _add_page_number_footer(section) -> None:
    """フッターに「- N -」を入れる。納品文書はページ番号が要る。"""
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("- ")
    _set_jp_font(run, size=9)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run2 = footer.add_run()
    _set_jp_font(run2, size=9)
    run2._element.append(fld_begin)
    run2._element.append(instr)
    run2._element.append(fld_end)

    run3 = footer.add_run(" -")
    _set_jp_font(run3, size=9)


def _add_header(section, title: str, doc_id: str) -> None:
    """全ページのヘッダーに「文書名 / 文書ID」を出す。納品文書の基本要件。"""
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run(f"{title}　{doc_id}")
    _set_jp_font(run, size=8)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def _fill_table(table, rows: list[tuple[str, ...]], header_row: bool = False) -> None:
    """(ラベル, 値) 形式の小さな表を埋める。ラベル列は網掛けする。"""
    for r_idx, values in enumerate(rows):
        for c_idx, text in enumerate(values):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            _add_inline(cell.paragraphs[0], text)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9.5)
                if header_row and r_idx == 0:
                    run.bold = True
            if (c_idx == 0 and not header_row) or (header_row and r_idx == 0):
                tc_pr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "D9E2F3")
                tc_pr.append(shd)


def _repo_relative(path: Path) -> str:
    """リポジトリ相対の表示名。リポジトリ外や相対パス指定でも落ちないようにする。"""
    try:
        return str(path.resolve().relative_to(ROOT))
    except ValueError:
        return str(path)


def _add_cover(doc: Document, doc_id: str, title: str, source: Path) -> None:
    """表紙・文書管理情報・承認欄。SIer の納品文書はこの3点が定型で入る。"""
    for _ in range(5):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(CONFIDENTIALITY)
    _set_jp_font(run, size=11)
    run.bold = True
    run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{PROJECT_NAME} システム開発")
    _set_jp_font(run, size=14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    _set_jp_font(run, size=26)
    run.bold = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(doc_id)
    _set_jp_font(run, size=12)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    for _ in range(7):
        doc.add_paragraph()

    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _fill_table(
        table,
        [
            ("文書ID", doc_id),
            ("版数", "1.0"),
            ("作成日", DOC_DATE),
            ("作成者", AUTHOR),
            ("機密区分", CONFIDENTIALITY),
            ("正本", _repo_relative(source)),
        ],
    )
    doc.add_page_break()

    # --- 2 ページ目: 文書管理情報と承認欄 ---
    heading = doc.add_heading(level=1)
    heading.text = ""
    _add_inline(heading, "文書管理情報")
    for run in heading.runs:
        _set_jp_font(run)
        run.font.color.rgb = RGBColor(0x1F, 0x37, 0x63)

    table = doc.add_table(rows=6, cols=2)
    table.style = "Table Grid"
    _fill_table(
        table,
        [
            ("文書名", title),
            ("文書ID", doc_id),
            ("プロジェクト名", PROJECT_NAME),
            ("配布先", "発注者、開発チーム、運用担当"),
            ("保管場所", "リポジトリ docs/sdlc/ 配下"),  # noqa: E501
            ("更新方法", "正本 Markdown を修正し scripts/build_delivery_docs.py で再生成"),
        ],
    )

    doc.add_paragraph()
    heading = doc.add_heading(level=1)
    heading.text = ""
    _add_inline(heading, "承認欄")
    for run in heading.runs:
        _set_jp_font(run)
        run.font.color.rgb = RGBColor(0x1F, 0x37, 0x63)

    table = doc.add_table(rows=3, cols=4)
    table.style = "Table Grid"
    _fill_table(
        table,
        [
            ("区分", "氏名", "承認日", "備考"),
            ("作成", "", "", ""),
            ("査閲", "", "", ""),
        ],
        header_row=True,
    )
    row = table.add_row()
    for c_idx, text in enumerate(("承認", "", "", "")):
        cell = row.cells[c_idx]
        cell.text = ""
        _add_inline(cell.paragraphs[0], text)
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(9.5)
        if c_idx == 0:
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "D9E2F3")
            tc_pr.append(shd)

    p = doc.add_paragraph()
    run = p.add_run("※ 承認欄は納品時に記入する。")
    _set_jp_font(run, size=9)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    doc.add_page_break()


def _add_toc(doc: Document) -> None:
    """目次フィールドを入れる。Word で開いたときにページ番号付きで展開される。"""
    heading = doc.add_heading(level=1)
    heading.text = ""
    _add_inline(heading, "目次")
    for run in heading.runs:
        _set_jp_font(run)
        run.font.color.rgb = RGBColor(0x1F, 0x37, 0x63)

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    _set_jp_font(run)

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")  # 開いた時に自動更新させる
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "（Word で開き F9 キーを押すと目次が生成されます）"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    for element in (begin, instr, separate, placeholder, end):
        run._element.append(element)

    doc.add_page_break()


def _add_figure_index(doc: Document, figures: list[str], tables: list[str]) -> None:
    """図目次・表目次。図表番号を振ったら索引も付けるのが納品文書の作法。"""
    if not figures and not tables:
        return

    if figures:
        heading = doc.add_heading(level=1)
        heading.text = ""
        _add_inline(heading, "図目次")
        for run in heading.runs:
            _set_jp_font(run)
            run.font.color.rgb = RGBColor(0x1F, 0x37, 0x63)
        for caption in figures:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(12)
            run = p.add_run(caption)
            _set_jp_font(run, size=10)

    if tables:
        heading = doc.add_heading(level=1)
        heading.text = ""
        _add_inline(heading, "表目次")
        for run in heading.runs:
            _set_jp_font(run)
            run.font.color.rgb = RGBColor(0x1F, 0x37, 0x63)
        for caption in tables:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(12)
            run = p.add_run(caption)
            _set_jp_font(run, size=10)

    doc.add_page_break()


def _add_caption(doc: Document, text: str) -> None:
    """図表キャプション。図は下、表は上に置く運用にする。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    _set_jp_font(run, size=9)
    run.bold = True
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)


def _clean_caption(text: str) -> str:
    """見出しからキャプション用の短い名前を作る。先頭の章番号と装飾を落とす。"""
    text = re.sub(r"^\d+(\.\d+)*\.?\s*", "", text)
    text = re.sub(r"[*`]", "", text)
    return text.strip() or "図表"


def _assign_figure_numbers(blocks: list[dict[str, Any]]) -> tuple[list[str], list[str]]:
    """図表に「図 N-M」「表 N-M」を割り当て、図目次・表目次用の一覧を返す。

    章番号は見出しに書かれた番号を優先する。文書側が「## 3. 方式設計」と採番して
    いるのに図だけ別採番になると、本文からの参照が食い違う。
    """
    chapter = 0
    fig_no = 0
    table_no = 0
    last_heading = "概要"
    figures: list[str] = []
    tables: list[str] = []

    for block in blocks:
        if block["type"] == "heading":
            if block["level"] == 2:
                if m := re.match(r"^(\d+)[\.\s]", block["text"].strip()):
                    chapter = int(m.group(1))
                else:
                    chapter += 1
                fig_no = 0
                table_no = 0
            if block["level"] >= 2:
                last_heading = block["text"]

        elif block["type"] == "code" and block["lang"] in _DIAGRAM_LANGS:
            fig_no += 1
            caption = f"図 {max(chapter, 1)}-{fig_no}　{_clean_caption(last_heading)}"
            block["caption"] = caption
            figures.append(caption)

        elif block["type"] == "table":
            table_no += 1
            caption = f"表 {max(chapter, 1)}-{table_no}　{_clean_caption(last_heading)}"
            block["caption"] = caption
            tables.append(caption)

    return figures, tables


def _insert_picture(doc: Document, png: Path) -> bool:
    """図を本文幅に収めて貼る。縦長の図は高さ基準にしてページからはみ出させない。"""
    max_width_in = 6.0
    max_height_in = 7.5
    try:
        from PIL import Image

        with Image.open(png) as img:
            width_px, height_px = img.size
        if width_px and height_px / width_px * max_width_in > max_height_in:
            doc.add_picture(str(png), height=Inches(max_height_in))
        else:
            doc.add_picture(str(png), width=Inches(max_width_in))
    except Exception:  # PIL 不在や壊れた PNG でも本文生成は止めない
        try:
            doc.add_picture(str(png), width=Inches(max_width_in))
        except Exception as exc:
            print(f"WARN: picture insert failed {png.name}: {exc}")
            return False
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    return True


def markdown_to_docx(md_path: Path, out_path: Path) -> tuple[int, int]:
    """Markdown 1 本を Word 1 本に変換する。戻り値は (図の数, 表の数)。"""
    text = md_path.read_text(encoding="utf-8")
    blocks = parse_markdown(text)
    figures, tables = _assign_figure_numbers(blocks)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(10.5)
    # Normal スタイルは rPr / rFonts を持たないことがあるので必ず作ってから設定する
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for attr in ("w:eastAsia", "w:ascii", "w:hAnsi"):
        rfonts.set(qn(attr), JP_FONT)

    title = next(
        (b["text"] for b in blocks if b["type"] == "heading" and b["level"] == 1), md_path.stem
    )
    title = re.sub(r"^WS2D-[A-Z]{2}-\d+[_\s]*", "", title).strip()
    doc_id_match = re.match(r"(WS2D-[A-Z]{2}-\d+)", md_path.stem)
    doc_id = doc_id_match.group(1) if doc_id_match else "—"

    _add_cover(doc, doc_id, title, md_path)
    _add_page_number_footer(doc.sections[0])
    _add_header(doc.sections[0], title, doc_id)
    _add_toc(doc)
    _add_figure_index(doc, figures, tables)

    rendered_figures = 0
    skipped_h1 = False
    for block in blocks:
        kind = block["type"]

        if kind == "heading":
            if block["level"] == 1 and not skipped_h1:
                skipped_h1 = True  # 表紙に出したので本文では省く
                continue
            heading = doc.add_heading(level=min(block["level"], 4))
            heading.text = ""
            _add_inline(heading, block["text"])
            for run in heading.runs:
                _set_jp_font(run)
                run.font.color.rgb = RGBColor(0x1F, 0x37, 0x63)

        elif kind == "para":
            _add_inline(doc.add_paragraph(), block["text"])

        elif kind == "list":
            for item in block["items"]:
                style_name = "List Number" if item["ordered"] else "List Bullet"
                p = doc.add_paragraph(style=style_name)
                if item["indent"]:
                    p.paragraph_format.left_indent = Pt(18 * (item["indent"] + 1))
                _add_inline(p, item["text"])

        elif kind == "table":
            _add_caption(doc, block.get("caption", ""))  # 表はキャプションが上
            _write_docx_table(doc, block)
            doc.add_paragraph()

        elif kind == "code":
            if block["lang"] in _DIAGRAM_LANGS:
                png = render_mermaid(block["text"])
                if png and _insert_picture(doc, png):
                    _add_caption(doc, block.get("caption", ""))  # 図はキャプションが下
                    rendered_figures += 1
                    continue
                note = doc.add_paragraph()
                run = note.add_run("［図の描画に失敗したため定義を掲載］")
                _set_jp_font(run, size=9)
                run.italic = True
                run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(12)
            run = p.add_run(block["text"])
            _set_jp_font(run, MONO_FONT, size=8.5)

        elif kind == "hr":
            doc.add_paragraph("─" * 40)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return rendered_figures, len(tables)


def _write_docx_table(doc: Document, block: dict[str, Any]) -> None:
    """表を Word のテーブルにする。列数はヘッダーに合わせて詰める。"""
    header = block["header"]
    rows = block["rows"]
    if not header:
        return
    cols = len(header)

    table = doc.add_table(rows=1, cols=cols)
    table.style = "Table Grid"
    table.autofit = True

    shading = "D9E2F3"
    for idx, name in enumerate(header):
        cell = table.rows[0].cells[idx]
        cell.text = ""
        _add_inline(cell.paragraphs[0], name)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), shading)
        tc_pr.append(shd)

    for raw in rows:
        cells = (raw + [""] * cols)[:cols]
        row = table.add_row()
        for idx, value in enumerate(cells):
            cell = row.cells[idx]
            cell.text = ""
            _add_inline(cell.paragraphs[0], value)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)


# ---------------------------------------------------------------- Excel 生成

_HEADER_FILL = PatternFill("solid", fgColor="1F3763")
_HEADER_FONT = Font(color="FFFFFF", bold=True, size=10, name=JP_FONT)
_BODY_FONT = Font(size=10, name=JP_FONT)


def _write_sheet(wb: Workbook, title: str, header: list[str], rows: list[list[Any]]) -> None:
    ws = wb.create_sheet(title=title[:31])
    ws.append(header)
    for cell in ws[1]:
        cell.fill = _HEADER_FILL
        cell.font = _HEADER_FONT
        cell.alignment = Alignment(vertical="center", horizontal="center", wrap_text=True)

    for row in rows:
        ws.append(row)

    for col_idx in range(1, len(header) + 1):
        letter = get_column_letter(col_idx)
        widest = max(
            [len(str(header[col_idx - 1]))]
            + [len(str(r[col_idx - 1])) for r in rows if col_idx - 1 < len(r)][:400]
            or [10]
        )
        ws.column_dimensions[letter].width = min(max(widest + 2, 10), 60)

    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.font = _BODY_FONT
            cell.alignment = Alignment(vertical="top", wrap_text=True)

    ws.freeze_panes = "A2"
    ws.auto_filter.ref = ws.dimensions


def _load(name: str) -> Any:
    path = ASBUILT_DIR / name
    if not path.exists():
        return []
    return json.loads(path.read_text(encoding="utf-8"))


def _sdlc_dir_for(doc_id: str) -> Path:
    """文書 ID に対応する正本 Markdown が置かれた工程ディレクトリを返す。

    Office ファイルは正本と同じ場所に並べる。別ディレクトリに分けると、
    文書を探すのに 2 か所見る必要が出て、更新漏れにも気づきにくい。
    """
    for path in SDLC_DIR.rglob("*.md"):
        if "_asbuilt" in path.parts:
            continue
        if path.stem.startswith(doc_id):
            return path.parent
    return SDLC_DIR / "70_delivery"


def build_api_workbook() -> Path:
    """エンドポイント一覧。200 行あるので Word の表では実用にならない。"""
    routes = _load("routes.json")
    wb = Workbook()
    wb.remove(wb.active)

    _write_sheet(
        wb,
        "APIエンドポイント一覧",
        ["No", "Blueprint", "メソッド", "パス", "エンドポイント名", "概要", "実装モジュール", "行"],
        [
            [
                i,
                r["blueprint"],
                "/".join(r["methods"]),
                r["path"],
                r["endpoint"],
                r["summary"],
                r["module"],
                r["line"],
            ]
            for i, r in enumerate(routes, 1)
        ],
    )

    counts: dict[str, int] = {}
    for r in routes:
        counts[r["blueprint"]] = counts.get(r["blueprint"], 0) + 1
    _write_sheet(
        wb,
        "Blueprint別集計",
        ["Blueprint", "エンドポイント数", "モジュール"],
        [
            [bp, n, next((r["module"] for r in routes if r["blueprint"] == bp), "")]
            for bp, n in sorted(counts.items(), key=lambda kv: -kv[1])
        ],
    )

    out = _sdlc_dir_for("WS2D-IF-001") / "WS2D-IF-001_APIエンドポイント一覧.xlsx"
    out.parent.mkdir(parents=True, exist_ok=True)
    wb.save(out)
    return out


def build_license_workbook() -> Path:
    """OSS ライセンス一覧。再配布判断に使うので集計シートを分ける。"""
    licenses = _load("licenses.json")
    req = (ROOT / "requirements.txt").read_text(encoding="utf-8").lower()
    req_dev_path = ROOT / "requirements-dev.txt"
    req_dev = req_dev_path.read_text(encoding="utf-8").lower() if req_dev_path.exists() else ""

    def usage(name: str) -> str:
        low = name.lower()
        if re.search(rf"^{re.escape(low)}[=<>~\[]", req, re.M):
            return "実行時依存"
        if re.search(rf"^{re.escape(low)}[=<>~\[]", req_dev, re.M):
            return "開発時のみ"
        return "推移的依存"

    wb = Workbook()
    wb.remove(wb.active)

    _write_sheet(
        wb,
        "OSSライセンス一覧",
        ["No", "パッケージ", "バージョン", "ライセンス", "用途区分", "エコシステム", "入手元"],
        [
            [
                i,
                lic["name"],
                lic.get("version", ""),
                lic.get("license", "UNKNOWN"),
                usage(lic["name"]),
                lic.get("ecosystem", "PyPI"),
                lic.get("homepage", ""),
            ]
            for i, lic in enumerate(licenses, 1)
        ],
    )

    tally: dict[str, int] = {}
    for lic in licenses:
        key = lic.get("license", "UNKNOWN") or "UNKNOWN"
        tally[key] = tally.get(key, 0) + 1

    copyleft = [
        lic for lic in licenses if re.search(r"GPL|AGPL|LGPL|MPL|EPL", lic.get("license", ""), re.I)
    ]
    _write_sheet(
        wb,
        "ライセンス種別集計",
        ["ライセンス", "件数", "コピーレフト該当"],
        [
            [k, v, "該当" if re.search(r"GPL|AGPL|LGPL|MPL|EPL", k, re.I) else ""]
            for k, v in sorted(tally.items(), key=lambda kv: -kv[1])
        ],
    )
    _write_sheet(
        wb,
        "要確認（コピーレフト・不明）",
        ["パッケージ", "バージョン", "ライセンス", "確認事項"],
        [
            [lic["name"], lic.get("version", ""), lic.get("license", ""), "再配布義務の確認"]
            for lic in copyleft
        ]
        + [
            [lic["name"], lic.get("version", ""), "UNKNOWN", "ライセンス特定が必要"]
            for lic in licenses
            if (lic.get("license") or "UNKNOWN") == "UNKNOWN"
        ],
    )

    out = _sdlc_dir_for("WS2D-LI-001") / "WS2D-LI-001_OSSライセンス一覧.xlsx"
    out.parent.mkdir(parents=True, exist_ok=True)
    wb.save(out)
    return out


def build_module_workbook() -> Path:
    """モジュール / クラス一覧。詳細設計のレビュー用。"""
    modules = _load("modules.json")
    wb = Workbook()
    wb.remove(wb.active)

    _write_sheet(
        wb,
        "モジュール一覧",
        ["No", "パス", "パッケージ", "概要", "LOC", "クラス数", "公開関数数", "内部依存"],
        [
            [
                i,
                m["path"],
                m["package"],
                m["summary"],
                m["loc"],
                len(m["classes"]),
                len(m["functions"]),
                ", ".join(m["internal_deps"]),
            ]
            for i, m in enumerate(modules, 1)
        ],
    )

    class_rows: list[list[Any]] = []
    for m in modules:
        for cls in m["classes"]:
            class_rows.append(
                [
                    m["path"],
                    cls["name"],
                    cls["summary"],
                    ", ".join(cls["bases"]),
                    len(cls["methods"]),
                    ", ".join(mm["name"] for mm in cls["methods"][:12]),
                ]
            )
    _write_sheet(
        wb,
        "クラス一覧",
        ["モジュール", "クラス名", "責務", "継承元", "公開メソッド数", "主要メソッド"],
        class_rows,
    )

    out = _sdlc_dir_for("WS2D-MD-001") / "WS2D-MD-001_モジュール・クラス一覧.xlsx"
    out.parent.mkdir(parents=True, exist_ok=True)
    wb.save(out)
    return out


def build_screen_workbook() -> Path:
    """画面（テンプレート）一覧。画面設計書のレビュー用。"""
    templates = _load("templates.json")
    routes = _load("routes.json")

    page_routes = [r for r in routes if "GET" in r["methods"] and not r["path"].startswith("/api")]
    wb = Workbook()
    wb.remove(wb.active)

    _write_sheet(
        wb,
        "画面ルート一覧",
        ["No", "パス", "エンドポイント", "Blueprint", "概要", "実装モジュール"],
        [
            [i, r["path"], r["endpoint"], r["blueprint"], r["summary"], r["module"]]
            for i, r in enumerate(page_routes, 1)
        ],
    )
    _write_sheet(
        wb,
        "テンプレート一覧",
        ["No", "テンプレート", "行数", "タイトル", "継承元", "include", "block"],
        [
            [
                i,
                t["path"],
                t["loc"],
                t["title"],
                ", ".join(t["extends"]),
                ", ".join(t["includes"]),
                ", ".join(t["blocks"]),
            ]
            for i, t in enumerate(templates, 1)
        ],
    )

    out = _sdlc_dir_for("WS2D-SD-001") / "WS2D-SD-001_画面・テンプレート一覧.xlsx"
    out.parent.mkdir(parents=True, exist_ok=True)
    wb.save(out)
    return out


def build_md_table_workbook(md_path: Path, out_name: str) -> Path | None:  # noqa: D401
    """表が主体の Markdown 文書を、表ごとのシートに展開する。"""
    blocks = parse_markdown(md_path.read_text(encoding="utf-8"))
    tables = [b for b in blocks if b["type"] == "table"]
    if not tables:
        return None

    # 各表の直前の見出しをシート名にする
    titles: list[str] = []
    current = "表"
    for block in blocks:
        if block["type"] == "heading":
            current = block["text"]
        elif block["type"] == "table":
            titles.append(current)

    wb = Workbook()
    wb.remove(wb.active)
    used: set[str] = set()
    for idx, (table, title) in enumerate(zip(tables, titles, strict=False), 1):
        name = re.sub(r"[\\/*?:\[\]]", "", title)[:28] or f"表{idx}"
        while name in used:
            name = f"{name[:26]}_{idx}"
        used.add(name)
        _write_sheet(wb, name, table["header"], table["rows"])

    out = md_path.parent / out_name
    out.parent.mkdir(parents=True, exist_ok=True)
    wb.save(out)
    return out


# ---------------------------------------------------------------- 実行


# 表が主体で、Excel でも配布したい文書
_TABLE_HEAVY = {
    "WS2D-TM-001": "WS2D-TM-001_トレーサビリティマトリクス.xlsx",
    "WS2D-DL-001": "WS2D-DL-001_不具合管理台帳.xlsx",
    "WS2D-TV-001": "WS2D-TV-001_テスト観点表.xlsx",
    "WS2D-NF-001": "WS2D-NF-001_非機能要件一覧.xlsx",
    "WS2D-RD-001": "WS2D-RD-001_要件一覧.xlsx",
}


def main() -> int:
    md_files = sorted(p for p in SDLC_DIR.rglob("*.md") if "_asbuilt" not in p.parts)

    word_count = 0
    total_figures = 0
    total_tables = 0
    for md_path in md_files:
        out_path = md_path.with_suffix(".docx")  # 正本と同じディレクトリに並べる
        try:
            figures, tables = markdown_to_docx(md_path, out_path)
            word_count += 1
            total_figures += figures
            total_tables += tables
        except Exception as exc:  # 1 本の失敗で全体を止めない
            print(f"WARN: docx failed {md_path.relative_to(SDLC_DIR)}: {exc}")

    excel_paths = [
        build_api_workbook(),
        build_license_workbook(),
        build_module_workbook(),
        build_screen_workbook(),
    ]

    for doc_id, out_name in _TABLE_HEAVY.items():
        match = next((p for p in md_files if p.stem.startswith(doc_id)), None)
        if match is None:
            print(f"WARN: {doc_id} の Markdown が見つからない")
            continue
        try:
            result = build_md_table_workbook(match, out_name)
            if result:
                excel_paths.append(result)
        except Exception as exc:
            print(f"WARN: xlsx failed {doc_id}: {exc}")

    print(f"Word  : {word_count} files -> {SDLC_DIR.relative_to(ROOT)}/<工程>/")
    print(f"        図 {total_figures} 点（PNG 埋め込み）/ 表 {total_tables} 点")
    print(f"Excel : {len(excel_paths)} files")
    for path in excel_paths:
        print(f"        {path.relative_to(ROOT)}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
